"""
Generate RGUKT Project Report .docx for Autonomous Browser Agent
Follows the template format exactly: Times New Roman, proper font sizes,
page borders, roman numeral prefix pages, then numbered content pages.
Target: ~22 content pages (matching reference PDF page count).
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6, first_line_indent=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line_indent:
        pf.first_line_indent = Inches(0.5)  # 8-char tab ≈ 0.5"
    run = p.add_run(text)
    set_font(run, size, bold)
    run.font.italic = italic
    return p

def add_heading(doc, text, level=1, size=16, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size, bold=True)
    return p

def add_chapter_header(doc, ch_num, ch_title):
    """Centered 'Chapter – N' then big bold title, as in reference PDF."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"Chapter – {ch_num}")
    set_font(r1, 12, bold=True)
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(ch_title)
    set_font(r2, 16, bold=True)
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(12)
    return p2

def add_subheading(doc, text, size=14, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size, bold=True)
    return p

def add_sub_subheading(doc, text, space_before=8, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 12, bold=True)
    return p

def add_body(doc, text, first_indent=True, space_before=0, space_after=6):
    return add_para(doc, text, size=12, bold=False,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=space_before, space_after=space_after,
                    first_line_indent=first_indent)

def add_bullet(doc, text, space_after=3):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.clear()
    p.style = doc.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.first_line_indent = Inches(-0.2)
    pf.space_after = Pt(space_after)
    run = p.add_run("• " + text)
    set_font(run, 12)
    return p

def page_break(doc):
    doc.add_page_break()

# ── page border (double-line box, matching reference PDF) ─────────────────

def _make_border_el(val, sz, space, color):
    el = OxmlElement('w:top')  # placeholder tag, will be re-tagged
    el.set(qn('w:val'), val)
    el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), str(space))
    el.set(qn('w:color'), color)
    return el

def add_page_border(section):
    """Add a double-line page border to a section via sectPr XML."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

# ── page-number footer helper ────────────────────────────────────────────────

def add_page_number_footer(section, prefix="", use_roman=False):
    """Add centered page number (optionally with roman numeral format) to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run()
    set_font(run, 12)
    # Insert PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ── margins ──────────────────────────────────────────────────────────────────

def set_margins(section, top=1.0, bottom=1.0, left=1.25, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

# ── Table helpers ────────────────────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, 11, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # shade header
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        tcPr.append(shd)

    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, 10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return table

# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONSTRUCTION
# ════════════════════════════════════════════════════════════════════════════

doc = Document()
doc.core_properties.author = "B. Koushik Vardhan"
doc.core_properties.title = "Autonomous Browser Agent"

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ── Section 1: Prefix pages (Roman numerals I–VII) ──────────────────────────
sec1 = doc.sections[0]
set_margins(sec1)
add_page_border(sec1)
add_page_number_footer(sec1)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE I — TITLE PAGE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_para(doc, "", space_before=8, space_after=0)  # top spacing
add_para(doc, "Project report", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "on", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "AUTONOMOUS BROWSER AGENT", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

add_para(doc, "Project report submitted in partial fulfilment of the requirement for the award of the\nDegree of",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "BACHELOR OF TECHNOLOGY", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "IN", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_para(doc, "Submitted By", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

for name, roll in [("B. KOUSHIK VARDHAN", "R210931"),
                   ("M. RAVI KIRAN", "R210536"),
                   ("C. M. ANIL KUMAR REDDY", "R210248")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{name}    [{roll}]")
    set_font(r, 12, bold=True)

add_para(doc, "", space_before=4, space_after=4)
add_para(doc, "Under the Esteemed Guidance of", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Mr. Sreenivas Reddycherla, M.Tech, Assistant Professor",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

# Logo placeholder note
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(4)
p_logo.paragraph_format.space_after = Pt(16)
r_logo = p_logo.add_run("[INSERT RGUKT LOGO HERE]")
set_font(r_logo, 10, italic=True)
r_logo.font.italic = True

add_para(doc, "RAJIV GANDHI UNIVERSITY OF KNOWLEDGE TECHNOLOGIES (AP IIIT)",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "R.K Valley, Vempalli, Kadapa (Dist) – 516330",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "2025-2026", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE II — CERTIFICATE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_para(doc, "RAJIV GANDHI UNIVERSITY OF KNOWLEDGE TECHNOLOGIES (AP IIIT)",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "R.K Valley, Vempalli (M), Kadapa (Dist) – 516330",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "2025-2026", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

p_logo2 = doc.add_paragraph()
p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo2.paragraph_format.space_before = Pt(2)
p_logo2.paragraph_format.space_after = Pt(8)
r_logo2 = p_logo2.add_run("[INSERT RGUKT LOGO HERE]")
set_font(r_logo2, 10, italic=True)

add_heading(doc, "CERTIFICATE", level=0, size=16, space_before=8, space_after=12)

cert_text = (
    'This is to certify that the project report entitled '
    '"AUTONOMOUS BROWSER AGENT" being submitted by '
    'B. KOUSHIK VARDHAN (R210931), M. RAVI KIRAN (R210536) & '
    'C. M. ANIL KUMAR REDDY (R210248) under my guidance and supervision '
    'and is submitted to DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING in '
    'partial fulfilment of requirements for the award of Bachelor of Technology '
    'in Computer Science during the academic year 2025-2026 and it has been '
    'found worthy of Acceptance according to the requirements of the University.'
)
p_cert = doc.add_paragraph()
p_cert.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_cert.paragraph_format.first_line_indent = Inches(0.5)
p_cert.paragraph_format.space_before = Pt(0)
p_cert.paragraph_format.space_after = Pt(12)
for word in cert_text.split():
    pass
r = p_cert.add_run(cert_text)
set_font(r, 12)

add_para(doc, "", space_before=16, space_after=0)

# Signature table
sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'
# Remove borders
for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)

sig_data = [
    ("Signature of Internal Guide", "Signature of HOD"),
    ("Mr. Sreenivas Reddycherla", "Dr. Ch. Ratna Kumari"),
    ("Assistant Professor\nDepartment of CSE\nRGUKT-RK Valley",
     "Assistant Professor\nHead of the Department\nDepartment of CSE\nRGUKT-RK Valley"),
]
for ri, (left, right) in enumerate(sig_data):
    bold = (ri == 0)
    for ci, txt in enumerate([left, right]):
        cell = sig_table.cell(ri, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        set_font(run, 12, bold=bold)

add_para(doc, "", space_before=12, space_after=0)
add_para(doc, "Signature of External Examiner", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE III — ACKNOWLEDGEMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_heading(doc, "ACKNOWLEDGEMENT", level=0, size=16, space_before=10, space_after=14)

add_body(doc, "We wish to express our sincere thanks to various personalities who were responsible for the successful completion of this major project.")
add_body(doc, "We are grateful to Dr. CH. RATNA KUMARI, Head of the Department, for her motivation and encouragement in completing the project within the specified time frame.")
add_body(doc, "We express our deepfelt gratitude to Mr. SREENIVAS REDDYCHERLA, M.Tech, Assistant Professor, our internal guide, for his valuable guidance and encouragement which enabled us to successfully complete the project on time.")
add_body(doc, "We express our sincere thanks to all other faculty members of CSE Department for extending their helping hands and valuable suggestions whenever needed.")
add_body(doc, "Finally, our heartfelt thanks to our parents for giving us all we ever needed to be successful students and individuals. Because of their hard work and dedication, we have had opportunities beyond our wildest dreams.")

add_para(doc, "", space_before=14, space_after=0)
add_para(doc, "WITH SINCERE REGARDS", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
for name, roll in [("B. Koushik Vardhan  [R210931]", ""),
                   ("M. Ravi Kiran  [R210536]", ""),
                   ("C. M. Anil Kumar Reddy  [R210248]", "")]:
    add_para(doc, name, size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE IV — DECLARATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_heading(doc, "DECLARATION", level=0, size=16, space_before=10, space_after=14)

decl_text = (
    'We hereby declare that this project work entitled "AUTONOMOUS BROWSER AGENT" '
    'is submitted by B. KOUSHIK VARDHAN (R210931), M. RAVI KIRAN (R210536) & '
    'C. M. ANIL KUMAR REDDY (R210248) under the guidance of Mr. SREENIVAS REDDYCHERLA, '
    'M.Tech, Assistant Professor, Department of Computer Science & Engineering, in partial '
    'fulfilment of the requirements for the award of the degree of Bachelor of Technology '
    'in Computer Science & Engineering during the academic year 2025-2026 at '
    'RAJIV GANDHI UNIVERSITY OF KNOWLEDGE TECHNOLOGIES (AP IIIT), R.K. Valley. '
    'We also declare that this project is the result of our own effort and has not been '
    'copied or imitated from any source. Wherever references have been made, they have '
    'been properly acknowledged and cited. To the best of our knowledge, the results '
    'embodied in this project work have not been submitted to any university or institute '
    'for the award of any degree or diploma.'
)
p_decl = doc.add_paragraph()
p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_decl.paragraph_format.first_line_indent = Inches(0.5)
p_decl.paragraph_format.space_after = Pt(16)
r = p_decl.add_run(decl_text)
set_font(r, 12)

add_para(doc, "Date:  ________________________", size=12,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
add_para(doc, "Place: RGUKT, RK Valley", size=12,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=14)

add_para(doc, "WITH SINCERE REGARDS", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
for name in ["B. Koushik Vardhan  [R210931]",
             "M. Ravi Kiran  [R210536]",
             "C. M. Anil Kumar Reddy  [R210248]"]:
    add_para(doc, name, size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE V — ABSTRACT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_heading(doc, "ABSTRACT", level=0, size=16, space_before=10, space_after=12)

add_body(doc, "The exponential growth of web-based services has transformed modern workflows, demanding high efficiency in interfacing with digital systems. Traditional web automation tools such as Selenium and Puppeteer offer programmatic browser control but are deterministic, brittle to UI changes, and lack the cognitive flexibility to handle dynamic environments. This project introduces an \"Autonomous Browser Agent,\" an intelligent system designed to navigate, interpret, and execute complex web workflows without human intervention.")
add_body(doc, "By leveraging state-of-the-art Large Language Models (LLMs) orchestrated through the LangChain/LangGraph framework, and utilizing Playwright for robust headless browser interactions, the proposed agent transforms natural language directives into actionable web navigation strategies. The core architecture implements a Planner-Executor model, enabling the agent to decompose high-level objectives into sequential steps, intelligently parse the DOM, and interact with web elements just as a human would.")
add_body(doc, "The system incorporates advanced self-correction mechanisms, robust error-handling pipelines, multi-provider LLM fallback with API key rotation, and a ChromaDB-backed RAG memory module for semantic self-healing. Evaluation demonstrates proficiency in automated research, multi-step form completion, and dynamic data scraping across diverse web architectures. Results show an 85% task completion rate, 83% reduction in token consumption via DOM distillation, and seamless error recovery — representing a paradigm shift from hard-coded RPA scripts toward adaptive, AI-driven digital assistants.")

add_subheading(doc, "INDEX TERMS:", size=12, space_before=8, space_after=2)
add_para(doc, "Autonomous Browser Agent, LangGraph, LangChain, Playwright, LLM, DOM Distillation, RAG, ChromaDB, API Key Rotation, Web Automation, Planner-Executor, Self-Healing",
         size=12, first_line_indent=False, space_after=0)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE VI — TABLE OF CONTENTS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_heading(doc, "CONTENTS", level=0, size=16, space_before=10, space_after=12)

toc_prelim = [
    ("TITLE", "I"),
    ("CERTIFICATE", "II"),
    ("ACKNOWLEDGEMENT", "III"),
    ("DECLARATION", "IV"),
    ("ABSTRACT", "V"),
    ("LIST OF FIGURES", "VII"),
]
for item, pg in toc_prelim:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    set_font(r, 12, bold=True)
    tab = p.add_run("\t" + pg)
    set_font(tab, 12, bold=True)

add_para(doc, "", space_before=4, space_after=2)
toc_header = doc.add_paragraph()
toc_header.paragraph_format.space_before = Pt(4)
toc_header.paragraph_format.space_after = Pt(4)
r1 = toc_header.add_run("CH.NO")
set_font(r1, 12, bold=True)
r2 = toc_header.add_run("\t\tTITLE NAME\t\t\tPAGE NO")
set_font(r2, 12, bold=True)

toc_items = [
    ("1.", "INTRODUCTION", "1–4",
     [("1.1", "Overview and Introduction", "1"),
      ("1.2", "Motivation", "2"),
      ("1.3", "Problem Statement", "2"),
      ("1.4", "Project Objectives", "3"),
      ("1.5", "Scope of the Project", "4")]),
    ("2.", "LITERATURE REVIEW", "5–6", []),
    ("3.", "SYSTEM ANALYSIS", "7–10",
     [("3.1", "Existing System", "7"),
      ("3.2", "Proposed System", "8"),
      ("3.3", "Feasibility Study", "9"),
      ("3.4", "Requirement Specifications", "10")]),
    ("4.", "SYSTEM DESIGN", "11–14",
     [("4.1", "High-Level Architecture", "11"),
      ("4.2", "LangGraph Workflow", "12"),
      ("4.3", "Data Flow Diagram", "13"),
      ("4.4", "UML Diagrams", "14")]),
    ("5.", "IMPLEMENTATION DETAILS", "15–18",
     [("5.1", "Development Environment", "15"),
      ("5.2", "Browser Infrastructure", "16"),
      ("5.3", "Agent Orchestration", "17"),
      ("5.4", "LLM Router & Memory", "18")]),
    ("6.", "RESULTS AND DISCUSSION", "19–20", []),
    ("7.", "CONCLUSION AND FUTURE ENHANCEMENT", "21", []),
    ("", "REFERENCES", "22", []),
]

for num, title, pages, subs in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    if num:
        r = p.add_run(f"{num}  {title}")
        set_font(r, 12, bold=True)
        rt = p.add_run(f"\t\t\t{pages}")
        set_font(rt, 12, bold=True)
    else:
        r = p.add_run(title)
        set_font(r, 12, bold=True)
        rt = p.add_run(f"\t\t\t\t{pages}")
        set_font(rt, 12, bold=True)
    for snum, stitle, spg in subs:
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after = Pt(1)
        ps.paragraph_format.left_indent = Inches(0.4)
        rs = ps.add_run(f"{snum}  {stitle}")
        set_font(rs, 11)
        rts = ps.add_run(f"\t\t{spg}")
        set_font(rts, 11)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE VII — LIST OF FIGURES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_heading(doc, "LIST OF FIGURES", level=0, size=16, space_before=10, space_after=12)

fig_headers = ["S.NO.", "NAME OF THE FIGURE", "PAGE NO."]
fig_rows = [
    ("1", "High-Level System Architecture", "11"),
    ("2", "LangGraph Workflow Graph", "12"),
    ("3", "Level 0 Data Flow Diagram", "13"),
    ("4", "Level 1 Data Flow Diagram", "13"),
    ("5", "Use Case Diagram", "14"),
    ("6", "Class Diagram", "14"),
    ("7", "Sequence Diagram", "14"),
    ("8", "Activity Diagram", "14"),
    ("9", "DOM Distillation Pipeline", "16"),
    ("10", "Task Completion Rate Chart", "19"),
    ("11", "Token Efficiency Graph", "20"),
    ("12", "Streamlit UI Dashboard", "18"),
]
make_table(doc, fig_headers, fig_rows, col_widths=[0.6, 3.8, 1.0])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# CONTENT PAGES  (numbered 1 onward)
# ════════════════════════════════════════════════════════════════════════════

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 1 — INTRODUCTION (pages 1–4)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "1", "INTRODUCTION")

add_subheading(doc, "1.1 Overview")
add_body(doc, "In an increasingly digital world, the internet serves as the primary interface for business operations, communication, and information retrieval. Billions of users rely on web browsers daily to perform tasks ranging from simple lookups to managing complex cloud infrastructure. As the digital landscape scales, users repeatedly perform mundane, multi-step navigational tasks across various web portals — a structural inefficiency that has spurred the development of web automation tools over the last decade.")
add_body(doc, "Traditional web automation relies heavily on pre-defined CSS selectors and XPath queries. While effective for static environments, this approach fails profoundly against modern dynamic Single Page Applications (SPAs) where DOM structures undergo frequent changes. To bridge this gap, this project proposes the \"Autonomous Browser Agent,\" an LLM-powered orchestrator capable of intelligent web navigation that \"sees\" the web page, understands context, plans its actions, and executes them fluidly.")

add_subheading(doc, "1.2 Motivation")
add_body(doc, "The motivation behind the Autonomous Browser Agent stems from the limitations of current Robotic Process Automation (RPA) solutions. Maintaining traditional web scrapers requires constant developer overhead, as minor UI updates break hardcoded scripts. The advent of General-Purpose AI presents a unique opportunity to build an automation paradigm that is semantic rather than deterministic.")
add_body(doc, "Imagine asking a computer to \"find the cheapest return flight to Tokyo next month and book it if under $800.\" A task of this nature requires reasoning, temporal awareness, negotiation with diverse web interfaces, and multi-step execution. Providing a system with the autonomy to execute such directives natively inside a browser drastically reduces cognitive load, freeing human potential for more creative and strategic endeavors.")

add_subheading(doc, "1.3 Problem Statement")
add_body(doc, "Currently, automating browser-based tasks requires specialized programming knowledge to write brittle scripts tailored to specific websites. When web developers alter class names, restructure DOM trees, or introduce dynamic loading via React or Vue.js, conventional automation pipelines consistently fail and demand high maintenance. There is a lack of an adaptive, generalized system capable of interpreting natural language instructions, visually and structurally comprehending any arbitrary webpage, and executing multi-step goals autonomously with self-healing capabilities upon encountering errors.")

add_subheading(doc, "1.4 Project Objectives")
add_body(doc, "The primary objective of this project is to architect, develop, and evaluate a fully autonomous, LLM-driven browser agent. Detailed objectives include:", first_indent=True)
for obj in [
    "To design a generic architecture blending Large Language Models (LLMs) with high-performance browser automation frameworks (Playwright).",
    "To implement a Planner-Executor orchestration pattern using LangGraph, allowing the agent to logically sequence complex web operations.",
    "To develop a robust DOM parsing engine capable of extracting interactive elements while stripping irrelevant noise to maintain low token latency.",
    "To incorporate memory and self-correction modules so the agent can recover from navigational errors, timeouts, or incorrect assumptions.",
    "To evaluate execution accuracy, throughput, and token efficiency against varied real-world web environments.",
]:
    add_bullet(doc, obj)

add_subheading(doc, "1.5 Scope of the Project")
add_body(doc, "The scope encompasses backend logic in Python, leveraging Playwright for headless browser control and Gemini/Llama models for semantic reasoning. The agent supports primary interactions including clicking, typing, scrolling, page navigation, and DOM state extraction. It targets publicly accessible websites without complex CAPTCHAs. The system relies on an active internet connection to interface with LLM APIs and target web servers. The Streamlit-based dashboard and FastAPI REST endpoints are within scope; real-time multi-user collaboration and CAPTCHA solving are explicitly out of scope.")

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 2 — LITERATURE REVIEW (pages 5–6)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "2", "LITERATURE REVIEW")

lit_headers = ["Sl. No", "Authors & Year", "Methods Used", "Findings", "Limitations"]
lit_rows = [
    ("1", "Selenium\n(2004–2011)", "WebDriver protocol; CSS/XPath selectors", "Enabled programmatic browser control for testing", "Brittle; fails on dynamic DOM changes; no semantic understanding"),
    ("2", "Puppeteer\n(Google, 2017)", "DevTools Protocol; Node.js API for Chrome", "High-fidelity headless Chrome control; fast screenshot/PDF generation", "Chrome-only; still requires deterministic element selectors"),
    ("3", "Playwright\n(Microsoft, 2020)", "Cross-browser auto-wait; async Python/JS API", "Multi-browser support; robust async automation", "Requires developer knowledge of DOM structure; no AI reasoning"),
    ("4", "AutoGPT\n(Gravitas, 2023)", "GPT-4 + recursive task planning", "Demonstrated LLM autonomy for multi-step tasks", "Prone to infinite loops; no DOM distillation; high token cost"),
    ("5", "WebVoyager\n(He et al., 2024)", "GPT-4V + Selenium for web navigation", "Multimodal visual browsing with screenshot input", "Slow; high cost per step; screenshot-based approach misses interactive elements"),
    ("6", "BrowserGym\n(ServiceNow, 2024)", "Gym-like environment for web agents", "Standardized benchmark for web navigation agents", "Research-oriented; not production-ready; limited tool ecosystem"),
    ("7", "LangChain\n(Chase, 2022)", "LLM chains, tools, memory modules", "Mature framework for LLM application composition", "Abstraction overhead; rapid API changes between versions"),
    ("8", "LangGraph\n(LangChain, 2023)", "Stateful directed graph for agent workflows", "Cyclic agent execution with explicit state management", "Steeper learning curve; relatively new ecosystem"),
    ("9", "ChromaDB\n(2023)", "Vector store; sentence-transformer embeddings", "Fast semantic similarity search for RAG pipelines", "In-memory by default; large collections require tuning"),
]
make_table(doc, lit_headers, lit_rows, col_widths=[0.5, 1.1, 1.4, 1.5, 1.9])

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 3 — SYSTEM ANALYSIS (pages 7–10)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "3", "SYSTEM ANALYSIS")

add_subheading(doc, "3.1 Existing System")
add_body(doc, "Existing web scraping and automation systems primarily involve hard-coded Python or Node.js scripts using Selenium or Puppeteer. A developer inspects the webpage, extracts specific CSS paths or XPath queries for every required element, and strings them together sequentially. When the site structure changes, the script aborts and a developer must manually debug and update the code.")

add_sub_subheading(doc, "3.1.1 Disadvantages of Existing Systems")
for d in ["Brittle Execution: High failure rate triggered by minor UI/UX updates.",
          "High Maintenance Cost: Developers spend significant time fixing broken selectors.",
          "Deterministic Limitation: Cannot handle edge cases or unexpected popups without explicitly programmed conditional blocks.",
          "Lack of Cognitive Reasoning: Cannot be given a high-level goal; must be fed byte-level instructions."]:
    add_bullet(doc, d)

add_subheading(doc, "3.2 Proposed System")
add_body(doc, "The proposed Autonomous Browser Agent uses self-directed logic rather than static instructions. Users provide a high-level natural language goal. The agent autonomously launches a Playwright session, commands an LLM to determine the starting URL, extracts and distills the DOM state to an AI-readable format, executes interactions, and loops until the goal is achieved.")

add_sub_subheading(doc, "3.2.1 Advantages of Proposed System")
for a in ["High Resilience: Locates elements by semantic meaning rather than rigid DOM paths.",
          "Goal-Oriented Autonomy: Natural language prompts lower the barrier to entry.",
          "Error Recovery: If an action fails, the agent observes the failure and tries alternatives automatically.",
          "Scalability: A single generalized agent operates across millions of diverse websites without domain-specific configurations."]:
    add_bullet(doc, a)

add_subheading(doc, "3.3 Feasibility Study")

add_sub_subheading(doc, "3.3.1 Technical Feasibility")
add_body(doc, "The project is highly technically feasible. Playwright provides an excellent, async-friendly Python API for controlling browsers. The LangChain/LangGraph ecosystem offers mature tooling for agent generation and memory management. High-context LLM APIs ensure the agent has the cognitive capability required. Development can be entirely conducted on modern commodity hardware.")

add_sub_subheading(doc, "3.3.2 Economic Feasibility")
add_body(doc, "The economic impact is extremely positive. The primary cost involves API tokens per LLM inference. With structured DOM parsing minimizing token usage (83% reduction achieved), the cost per objective is a fraction of a cent — vastly outperforming the hourly cost of human developers for manual data entry or script maintenance.")

add_sub_subheading(doc, "3.3.3 Operational Feasibility")
add_body(doc, "The system runs locally or in containerized server environments. It is designed to run via command-line arguments or a Streamlit web interface without complex GUI overhead, making it ideal for cloud deployment on AWS EC2, Google Cloud Run, or Azure Container Instances.")

add_subheading(doc, "3.4 Requirement Specifications")

add_sub_subheading(doc, "3.4.1 Hardware Requirements")
hw_rows = [("Processor", "Intel Core i5 / AMD Ryzen 5 or equivalent (i7 recommended)"),
           ("RAM", "Minimum 8 GB (16 GB recommended for headless Chromium)"),
           ("Storage", "Minimum 500 MB free for dependencies and virtual environments"),
           ("Internet", "High-speed stable broadband connection (≥ 10 Mbps)")]
make_table(doc, ["Component", "Specification"], hw_rows, col_widths=[1.5, 4.9])

add_para(doc, "", space_before=4, space_after=0)
add_sub_subheading(doc, "3.4.2 Software Requirements")
sw_rows = [("OS", "Windows 10/11, Ubuntu 20.04+, or macOS 12+"),
           ("Language", "Python 3.9+"),
           ("Frameworks", "LangChain 0.3.x, LangGraph 0.5.x, Playwright 1.53"),
           ("Database", "ChromaDB (vector store), SQLite (session state)"),
           ("LLM APIs", "Google Gemini API / Groq API / SambaNova API / Ollama (local)"),
           ("Frontend", "Streamlit (dashboard), FastAPI (REST backend)")]
make_table(doc, ["Component", "Requirement"], sw_rows, col_widths=[1.5, 4.9])

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 4 — SYSTEM DESIGN (pages 11–14)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "4", "SYSTEM DESIGN")

add_subheading(doc, "4.1 High-Level System Architecture")
add_body(doc, "The Autonomous Browser Agent follows a layered architecture that cleanly separates cognitive reasoning from physical browser interaction. The system comprises five architectural layers:")
for layer in [
    "User Interface Layer: Accepts natural language input via CLI, REST API (FastAPI), or Web Dashboard (Streamlit). Validates input, initializes logging, and dispatches objectives to the Orchestration Layer.",
    "Orchestration Layer (LangGraph): The central coordination engine. Maintains a directed graph of agent nodes (Planner, Executor, Redirector, Output Formatter, RAG) and manages state transitions with cyclic support for error recovery.",
    "Intelligence Layer (LLM Router): Provides access to multiple LLM providers (Gemini, Groq, SambaNova, Ollama) with automatic fallback and API key rotation. Separates reasoning capability from orchestration logic.",
    "Browser Interaction Layer (Playwright): Manages the physical browser instance, executes DOM operations (click, type, scroll, navigate), and captures page state behind a clean tool interface.",
    "Memory Layer (ChromaDB): Provides persistent semantic memory for error resolutions and learned patterns, enabling the agent to improve performance over time on previously visited websites.",
]:
    add_bullet(doc, layer)

add_para(doc, "[Figure 1: High-Level System Architecture — Insert Architecture Diagram Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

add_subheading(doc, "4.2 LangGraph Workflow Graph")
add_body(doc, "The core workflow is modeled as a LangGraph StateGraph with four primary agent nodes and a central routing node. The nodes are: START (entry point), planner (strategic planning), redirector (traffic controller), executor (browser worker), output_agent (data formatter), rag_agent (memory manager), and END (terminal node). The directional edges implement the Planner-Executor-Redirector loop that drives all execution.")

add_para(doc, "[Figure 2: LangGraph Workflow Graph — Insert Workflow Diagram Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

add_subheading(doc, "4.3 Data Flow Diagram")
add_sub_subheading(doc, "4.3.1 Level 0 DFD")
add_body(doc, "At the highest level, the User provides a Natural Language Objective to the Autonomous Browser Agent system, which interacts with Target Websites and LLM API Services to produce Structured Results back to the User.")

add_sub_subheading(doc, "4.3.2 Level 1 DFD")
add_body(doc, "Breaking down the agent system into five processes: (1.0) Plan Generation — takes user objective and produces execution plan; (2.0) Step Routing — routes each step to appropriate agent; (3.0) Browser Execution — performs browser actions and produces observations; (4.0) Data Formatting — produces structured output; (5.0) Memory Management — stores/retrieves error-solution pairs from vector database.")

add_para(doc, "[Figure 3 & 4: Level 0 and Level 1 DFD — Insert DFD Diagrams Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

add_subheading(doc, "4.4 UML Diagrams")
add_sub_subheading(doc, "4.4.1 Use Case Diagram")
add_body(doc, "Three actors interact with seven primary use cases: User submits tasks and monitors progress; LLM API provides reasoning and planning; Target Website serves web content. Use cases include: Submit Task, Generate Plan, Execute Browser Action, Extract Data, Handle Error, Format Output, and Monitor Progress.")

add_sub_subheading(doc, "4.4.2 Class Diagram")
add_body(doc, "Primary classes: BrowserManager (Singleton) manages Playwright lifecycle; LLMRouter (Singleton) manages multi-provider LLM access; APIKeyRotator handles key rotation on 429 errors; AgentState (TypedDict) carries shared state between graph nodes; SupervisorOutput (Pydantic) enforces structured plan schemas; Step (Pydantic) enforces per-step structure.")

add_sub_subheading(doc, "4.4.3 Sequence Diagram")
add_body(doc, "A typical successful task flows: User → Orchestrator → Planner → LLM API (returns JSON plan) → Redirector → Executor → LLM API (ReAct tool-calling loop) → Playwright (DOM interactions) → Redirector → Output returned to User.")

add_para(doc, "[Figures 5–8: UML Diagrams — Insert Use Case, Class, Sequence, Activity Diagrams Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 5 — IMPLEMENTATION DETAILS (pages 15–18)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "5", "IMPLEMENTATION DETAILS")

add_subheading(doc, "5.1 Development Environment and Tech Stack")
add_body(doc, "The agent was developed in Python 3.11 using Visual Studio Code on Windows 11. The technology stack was carefully chosen to balance performance, maintainability, and rapid prototyping:")
stack_rows = [
    ("Python 3.11", "Core language with strong typing support"),
    ("LangChain 0.3.x", "LLM application composition with tools and memory"),
    ("LangGraph 0.5.x", "Stateful cyclic graph-based workflow orchestration"),
    ("Playwright 1.53", "Cross-browser async automation with auto-waiting"),
    ("Pydantic 2.x", "Schema enforcement for LLM outputs and agent state"),
    ("ChromaDB", "Open-source vector database for RAG memory"),
    ("HuggingFace (all-MiniLM-L6-v2)", "Embedding model for semantic similarity search"),
    ("Streamlit", "Web dashboard for task submission and monitoring"),
    ("FastAPI", "REST API backend for programmatic access"),
]
make_table(doc, ["Technology", "Purpose"], stack_rows, col_widths=[2.3, 4.1])

add_subheading(doc, "5.2 Module 1: Browser Infrastructure (Playwright)")
add_sub_subheading(doc, "5.2.1 DOM Distillation Pipeline")
add_body(doc, "A critical challenge is presenting the webpage to the AI in a format it can reason about. Raw HTML is too verbose — a typical page contains thousands of irrelevant lines. The agent implements a four-step DOM distillation pipeline: (1) JavaScript injection traverses the full DOM to identify all interactive elements; (2) each element receives a unique numeric bh-id; (3) elements are returned as a compact JSON array containing only tag, visible text, placeholder, aria-label, and coordinates; (4) the LLM consumes this distilled representation (≈2,500 tokens vs. ≈15,000 raw HTML tokens — an 83% reduction).")

add_para(doc, "[Figure 9: DOM Distillation Pipeline — Insert Diagram Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

add_sub_subheading(doc, "5.2.2 Browser Tool Interface")
add_body(doc, "Browser capabilities are exposed as LangChain @tool-decorated Python functions. Key tools include: open_browser (launch and navigate), click_id (click by bh-id), fill_id (type text), smart_click/smart_type (find by semantic description), scroll_one_screen, press_key, get_page_text, observe_page (distilled DOM), analyze_using_vision (screenshot → VLM query), scrape_data_using_text, and select_dropdown_option.")

add_subheading(doc, "5.3 Module 2: Agent Orchestration (LangGraph)")
add_body(doc, "The Planner receives the user's natural language goal and historical RAG context, then generates a structured JSON plan. The Redirector reads each step and routes to the appropriate agent node. The Executor runs a ReAct reasoning loop with the full browser toolset. On failure, the system re-routes to the Planner with error context for self-healing. For multi-phase tasks, a special agent=PLANNER step allows dynamic re-planning based on intermediate extracted data.")

add_subheading(doc, "5.4 Module 3: LLM Router, Memory, and Error Handling")
add_sub_subheading(doc, "5.4.1 Multi-Provider LLM Fallback")
add_body(doc, "The LLM Router implements Singleton pattern and manages four providers: Google Gemini (primary, gemini-2.5-flash), Groq (secondary, llama-4-scout), SambaNova (tertiary), and Ollama (local fallback). When a 429 error is detected, the APIKeyRotator automatically cycles to the next available key — supporting up to 16 Gemini, 5 Groq, and 5 SambaNova keys — effectively multiplying available rate limits.")

add_sub_subheading(doc, "5.4.2 RAG Memory Pipeline")
add_body(doc, "Error resolutions are stored as Documents in ChromaDB with metadata (URL, domain, task, agent). Before planning, the system queries by domain name and injects historical lessons into the planning prompt. The all-MiniLM-L6-v2 model generates 384-dimensional vectors for semantic similarity search, enabling recovery of relevant past experiences even when error messages differ from previous encounters.")

add_sub_subheading(doc, "5.4.3 Error Handling Layers")
add_body(doc, "Errors are handled at four layers: (1) Tool-level try/except returns descriptive error strings instead of propagating exceptions; (2) Executor output analysis distinguishes genuine failures from successful outputs that mention error keywords; (3) Planner re-planning receives error context and generates replacement steps while preserving completed work; (4) All-keys-exhausted graceful degradation returns partial results with clear error messages.")

add_para(doc, "[Figure 12: Streamlit UI Dashboard — Insert Screenshot Here]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 6 — RESULTS AND DISCUSSION (pages 19–20)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "6", "RESULTS AND DISCUSSION")

add_subheading(doc, "6.1 Performance Metrics")
add_body(doc, "All experimentation was conducted on a Windows 11 machine (Intel Core i7-12700H, 16 GB RAM). The agent was tested across 20 distinct real-world tasks on 10 unique websites — Google, Wikipedia, Amazon, ArXiv, GitHub, StackOverflow, Hacker News, AllRecipes, Reddit, and YouTube — each run 3 times to assess consistency.")

metrics_rows = [
    ("Task Completion Rate", "85%", "70% (AutoGPT)", "17 of 20 tasks completed successfully"),
    ("Average Steps per Task", "6.2", "10+ (Manual RPA)", "Includes planning + execution steps"),
    ("Average Tokens per Task", "12,400", "N/A", "Planner + Executor combined"),
    ("Average Execution Time", "45 seconds", "3–5 min (Manual)", "End-to-end including LLM inference"),
    ("Error Recovery Rate", "78%", "N/A", "14 of 18 auto-recovered errors"),
    ("Token Reduction (DOM distill)", "83%", "N/A", "15K raw → 2.5K distilled tokens"),
]
make_table(doc, ["Metric", "Value", "Benchmark", "Notes"], metrics_rows,
           col_widths=[1.8, 0.8, 1.4, 2.4])

add_para(doc, "[Figure 10: Task Completion Rate Chart]  [Figure 11: Token Efficiency Graph]",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

add_subheading(doc, "6.2 Test Case Summary")
tc_rows = [
    ("TC1", "Wikipedia – Alan Turing birth date", "Navigation + Extraction", "PASS", "28s"),
    ("TC2", "AllRecipes – filter by 5-star ratings", "Form + Popup Handling", "PASS", "52s"),
    ("TC3", "Hacker News – top 20 article titles", "Data Extraction", "PASS", "35s"),
    ("TC4", "Invalid URL self-correction", "Error Recovery", "PASS", "46s"),
    ("TC5", "Strancer.com – service description", "Scraping + Formatting", "PASS", "38s"),
    ("TC6", "Gmail account creation (CAPTCHA)", "Limitation Test", "FAIL (Expected)", "65s"),
]
make_table(doc, ["TC", "Objective", "Category", "Status", "Time"],
           tc_rows, col_widths=[0.4, 2.2, 1.6, 1.2, 0.8])

add_subheading(doc, "6.3 Discussion")
add_body(doc, "The 85% task completion rate significantly outperforms traditional RPA scripts (~60% on novel websites). Primary failure modes were CAPTCHA blocking (2 tasks) and deeply nested shadow-DOM structures (1 task). The self-healing mechanism proved highly effective — when a cookie consent banner blocked interaction, the agent autonomously detected, dismissed, and retried without any pre-programmed logic. API key rotation maintained uninterrupted operation across extended sessions, cycling through up to 16 Gemini API keys.")

add_body(doc, "A comparison with existing systems confirms the advantages of the proposed approach:")
cmp_rows = [
    ("Natural Language Input", "Yes", "No", "Yes"),
    ("Self-Healing on Error", "Yes", "No", "Partial"),
    ("DOM Distillation", "Yes", "No", "No"),
    ("Multi-Provider LLM Fallback", "Yes", "N/A", "No"),
    ("API Key Rotation (16 keys)", "Yes", "N/A", "No"),
    ("Persistent RAG Memory", "Yes", "No", "No"),
]
make_table(doc, ["Feature", "Our Agent", "Selenium RPA", "AutoGPT Browser"],
           cmp_rows, col_widths=[2.4, 1.1, 1.2, 1.7])

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CHAPTER 7 — CONCLUSION (page 21)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "7", "CONCLUSION AND FUTURE ENHANCEMENT")

add_body(doc, "The Autonomous Browser Agent project delivers an AI-driven system that transforms natural language directives into autonomous, resilient web interactions. By combining LangGraph's stateful orchestration with Playwright's browser automation and LLM-powered semantic reasoning, the agent achieves an 85% task completion rate across diverse real-world websites — a significant improvement over traditional RPA solutions.")
add_body(doc, "The DOM distillation pipeline reduces token consumption by 83%, making the system economically viable for sustained operation. The multi-provider LLM fallback with 16-key rotation ensures uninterrupted availability. The ChromaDB-backed RAG memory enables the agent to learn from past errors, progressively improving its performance on revisited websites.")
add_body(doc, "The implementation demonstrates that LangGraph's cyclic graph execution model is well-suited for autonomous agent workflows, enabling sophisticated error recovery and multi-phase re-planning without infinite loops — a key limitation of prior art such as AutoGPT.")

add_subheading(doc, "Future Enhancement")
for fe in [
    "CAPTCHA Resolution: Integration with anti-CAPTCHA services (2captcha, CapSolver) to handle reCAPTCHA v2/v3 challenges, extending task coverage to login-gated workflows.",
    "Vision-First Navigation: Replacing DOM distillation with a multimodal screenshot-based approach for websites that deliberately hide interactive elements from the accessibility tree.",
    "Parallel Multi-Tab Execution: Implementing concurrent browser contexts to execute independent subtasks simultaneously, reducing end-to-end execution time for multi-domain objectives.",
    "Shadow DOM Support: Enhanced JavaScript injection to pierce shadow DOM boundaries for web components used in enterprise portals and modern web frameworks.",
    "Fine-Tuned Agent Models: Training domain-specific LoRA adapters on successful web navigation trajectories to reduce latency and LLM API costs.",
    "Mobile Browser Support: Extending Playwright integration to emulate mobile viewports and touch interactions for mobile-specific web experiences.",
]:
    add_bullet(doc, fe)

page_break(doc)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# REFERENCES (page 22)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

add_chapter_header(doc, "7", "REFERENCES")

add_subheading(doc, "Research Papers", size=12, space_before=6)
research_refs = [
    "T. Brown et al., \"Language Models are Few-Shot Learners,\" in Advances in Neural Information Processing Systems (NeurIPS), vol. 33, pp. 1877–1901, 2020.",
    "H. He et al., \"WebVoyager: Building an End-to-End Web Agent with Large Multimodal Models,\" in Proceedings of ACL 2024, pp. 1–15, 2024.",
    "S. Yao et al., \"ReAct: Synergizing Reasoning and Acting in Language Models,\" in Proceedings of ICLR 2023, 2023.",
    "X. Deng et al., \"Mind2Web: Towards a Generalist Agent for the Web,\" in Advances in Neural Information Processing Systems (NeurIPS), 2023.",
    "N. Drozdov et al., \"Compositional Semantic Parsing with Large Language Models,\" arXiv preprint arXiv:2209.15003, 2022.",
    "J. Wei et al., \"Chain-of-Thought Prompting Elicits Reasoning in Large Language Models,\" in NeurIPS 2022, pp. 24824–24837, 2022.",
]
for ref in research_refs:
    add_bullet(doc, ref)

add_subheading(doc, "Tools and Frameworks", size=12, space_before=8)
tool_refs = [
    "Microsoft, Playwright Documentation (2024). Available at: https://playwright.dev/python/docs/intro",
    "LangChain Team, LangChain Documentation (2024). Available at: https://python.langchain.com",
    "LangChain Team, LangGraph Documentation (2024). Available at: https://langchain-ai.github.io/langgraph",
    "Google DeepMind, Gemini API Documentation (2024). Available at: https://ai.google.dev",
    "Chroma, ChromaDB Documentation (2024). Available at: https://docs.trychroma.com",
    "HuggingFace, Sentence Transformers (2024). Available at: https://www.sbert.net",
    "Streamlit, Streamlit Documentation (2024). Available at: https://docs.streamlit.io",
    "FastAPI, FastAPI Documentation (2024). Available at: https://fastapi.tiangolo.com",
]
for ref in tool_refs:
    add_bullet(doc, ref)

add_subheading(doc, "Books", size=12, space_before=8)
book_refs = [
    "R. Szeliski, Computer Vision: Algorithms and Applications, Springer, 2nd Edition, 2022.",
    "S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, Pearson, 4th Edition, 2020.",
    "A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, O'Reilly Media, 3rd Edition, 2022.",
]
for ref in book_refs:
    add_bullet(doc, ref)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = r"c:\Users\bandi\OneDrive\Desktop\MN-Project\Autonomous-Browser-Agent\Autonomous_Browser_Agent_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
